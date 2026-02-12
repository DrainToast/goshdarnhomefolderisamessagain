import torch
from transformers import PreTrainedTokenizerBase
from sentence_transformers import SentenceTransformer
from typing import cast
from datetime import datetime
import filetype as ft  # type: ignore
from docx import Document as Document_docx
from keybert import KeyBERT # type: ignore
from keyphrase_vectorizers import KeyphraseCountVectorizer # type: ignore
from pathlib import Path
from pydantic import BaseModel
import emoji
import io
import regex as re
import unicodedata

PATH: str = "~/Documents/dominican friars data/"

class Metadata(BaseModel):
    title: str
    author: str
    keywords: list[str]
    categories: list[str]
    created: datetime | None
    modified: datetime | None
    last_modified_by: str

class _Document(BaseModel):
    path: Path
    raw: bytes
    filetype: str
    metadata: Metadata | None = None
    text: str | None = None
    keywords: list[str] | None = None
    categories: list[str] | None = None
    summary: list[str] | None = None


class BaseClass():
    def __init__(
        self,
        path_in: str
    ) -> None:
        path = Path(path_in).expanduser()
        if not path.exists():
            raise FileNotFoundError(f"The path supplied does not exist: \"{path_in}\".")

        self.docs: list[_Document] = []
        for doc_path in Path(path).expanduser().rglob("*"):
            if doc_path.is_file():
                with open(doc_path, "rb") as f:
                    doc_raw = f.read()
                    if filetype := ft.guess(doc_raw):
                        doc_filetype: str = filetype.extension
                    else:
                        doc_filetype = "Unknown"

                self.docs.append(
                    _Document(
                        path=doc_path,
                        raw=doc_raw,
                        filetype=doc_filetype
                    )
                )
            else:
                pass


        self._device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")
        if self._device.type == "cuda":
            torch.cuda.set_device(self._device)
            torch.cuda.empty_cache()
            torch.backends.cudnn.benchmark = True
            torch.backends.cuda.matmul.allow_tf32 = True
            torch.backends.cudnn.allow_tf32 = True

        self._keywords_model_initialized = False
        self._categories_model_initialized = False

        self._vectorizer: KeyphraseCountVectorizer
        self._keyword_model: KeyBERT
        self._categories_model: SentenceTransformer
        self._categories_tokenizer: PreTrainedTokenizerBase

        self._entailment_id: int = 2

        self._seed_keywords: list[str] = [""]
        self._categories: list[str] = [""]
        self._categories_embeddings: torch.Tensor


    def extract_document_content(self, doc: _Document) -> None:
        if doc.filetype == "docx":
            docx_doc = Document_docx(io.BytesIO(doc.raw))
            props = docx_doc.core_properties
            doc.metadata = Metadata(
                title=props.title,
                author=props.author,
                keywords=[props.keywords] if props.keywords else [],
                categories=[props.category] if props.category else [],
                created=props.created,
                modified=props.modified,
                last_modified_by=props.last_modified_by,
            )

            extracted = "\n".join(paragraph.text for paragraph in docx_doc.paragraphs)
            extracted = emoji.demojize(extracted)
            extracted = unicodedata.normalize("NFKD", extracted)
            extracted = re.sub(r'https?://\S+|www\.\S+', '', extracted)
            extracted = re.sub(r'\b[\w.+-]+@[\w-]+\.[\w.-]+\b', '', extracted)
            extracted = re.sub(r'[@#]\w+', '', extracted)
            extracted = re.sub(r'[ \t\r\f\v]+', ' ', extracted)
            extracted = re.sub(r'(?<!\n)\n\n(?!\n)', '\n', extracted)
            extracted = re.sub(r'\n{3,}', '\n\n', extracted)

            doc.text = extracted

        else:
            raise Exception(f"Content exctraction for filetype \"{doc.filetype}\", of path \"{doc.path}\".")


    def get_document_keywords(self, doc: _Document, score: float) -> None:
        if not self._keywords_model_initialized:
            self._vectorizer =  KeyphraseCountVectorizer(
                pos_pattern=r'<J.*>*<N.*>+|<N.*>+|<PROPN.*>+',
                stop_words='english',
                lowercase=True,
                spacy_pipeline='en_core_web_trf',
            )
            self._keyword_model = KeyBERT(model="all-mpnet-base-v2")

            if Path("seed_keywords.txt").exists():
                self._seed_keywords = [
                    line.strip()
                    for line
                    in Path("seed_keywords.txt").read_text(encoding="utf-8").splitlines()
                    if line.strip()
                ]

        if not doc.text:
            raise Exception(f"\"text\" has not been extracted yet from {doc.path}.")
        raw_keywords = self._keyword_model.extract_keywords(
            doc.text,
            vectorizer=self._vectorizer,  # type: ignore[arg-type]
            seed_keywords=self._seed_keywords,
            use_mmr=True,
            diversity=0.65
        )

        doc.keywords = [
            kw
            for kw, prob
            in cast(list[tuple[str, float]], raw_keywords)
            if prob >= score
        ]

    def get_document_categories(self, doc: _Document) -> None:
        if not doc.text or not doc.text.strip():
            raise ValueError("Document text is empty")

        if not self._categories_model_initialized:
            self._categories_model = SentenceTransformer("BAAI/bge-base-en-v1.5")

            if Path("categories.txt").exists():
                self._categories = [
                    line.strip()
                    for line
                    in Path("categories.txt").read_text(encoding="utf-8").splitlines()
                    if line.strip()
                ]
            else:
                raise FileNotFoundError("categories.txt not found")


            self._categories_embeddings = self._categories_model.encode(
                self._categories,
                convert_to_tensor=True,
                batch_size=64,
                device=[self._device]
            )

            self._categories_model_initialized = True

        doc_embedding = self._categories_model.encode(
            doc.text,
            convert_to_tensor=True,
            device=[self._device]
        )

        cos_scores = torch.nn.functional.cosine_similarity(
            doc_embedding.unsqueeze(0),
            self._categories_embeddings,
            dim=1
        )

        top_k = 200
        top_probs, top_indices = torch.topk((cos_scores + 1) / 2, k=top_k)

        top_categories = [
            {"label": self._categories[int(idx)], "score": float(score)}
            for idx, score in zip(top_indices, top_probs)
        ]

        for i in top_categories:
            print(i)



if __name__ == "__main__":
    bc = BaseClass(PATH)
    doc = bc.docs[11]
    bc.extract_document_content(doc)
    # bc.get_document_keywords(doc, 0.25)
    bc.get_document_categories(doc)


